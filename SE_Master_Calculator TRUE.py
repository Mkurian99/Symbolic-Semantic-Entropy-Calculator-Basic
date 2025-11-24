"""
================================================================================
SYMBOLIC ENTROPY (SE) MASTER CALCULATOR - Complete Implementation
================================================================================

FEATURES:
✓ Layer 1: Core SE Analysis (H + Σ with proper KL divergence)
✓ Layer 2: Peak/Valley Analysis with Text Quotes
✓ Multi-word Phrase Support (semantic token definition)
✓ Multi-format Support (.txt and .docx files)

THEORETICAL FOUNDATION:
    SE = (H, Σ)
    
    Where:
        H = Shannon entropy (bits/semantic-token) - lexical unpredictability
        Σ = KL divergence (bits/semantic-token) - motif clustering beyond baseline

KEY INNOVATION - Semantic Token Definition:
    - "One Ring" = 1 semantic token (merged pre-processing)
    - Phrases in motif dictionary are pre-merged: "one ring" → "one_ring"
    - Results in bits per SEMANTIC TOKEN, not word-token
    - Captures linguistic compression naturally

MOTIF DICTIONARY FORMAT:
    motif_dict = {
        'Category Name': {
            'phrases': ['multi word phrase', 'another phrase'],  # Optional
            'words': ['single', 'word', 'tokens']                # Optional
        }
    }
    
    - Phrase-only motifs: Only include 'phrases' key
    - Word-only motifs: Only include 'words' key  
    - Mixed motifs: Include both keys

IMPLEMENTATION STANDARDS:
    - Adaptive window sizing (auto-scales to text length)
    - Default: ~110 windows with 50% overlap
    - To adjust granularity: Change TARGET_WINDOWS (line 108)
      - Higher value = more windows = finer resolution
      - Lower value = fewer windows = coarser resolution
    - Global baseline (π_k) from full text distribution
    - Whole-word/phrase matching with word boundaries
    - Falsifiable via shuffle testing (Σ → 0 when randomized)

OUTPUTS:
    - <textname>_se_heatmap.png          (dual heatmap: raw density + KL)
    - <textname>_se_timeseries.png       (H and Σ line graphs)
    - <textname>_peaks_valleys.png       (top 3 peaks/valleys with quotes)
    - <textname>_se_results.csv          (full numerical results)
    - <textname>_peaks_valleys_text.csv  (detailed peak/valley excerpts)

DEPENDENCIES:
    - numpy, pandas, matplotlib, scipy
    - python-docx (for .docx support: pip install python-docx)

USAGE:
    python SE_Master_Calculator.py <path_to_file>
    
    Or edit TEXT_FILE variable (line 107)
    
    To adjust window granularity:
    - Edit TARGET_WINDOWS (line 108)
    - Default: 110 windows
    - Higher = finer resolution (more windows, smaller size)
    - Lower = coarser resolution (fewer windows, larger size)

VERSION: 2.0.0 - Multi-word phrase support
AUTHOR: Kurian, M. (2025)
================================================================================
"""

import re
import os
import sys
import numpy as np
import pandas as pd
import random
import matplotlib.pyplot as plt
from collections import Counter
from scipy.signal import find_peaks

# Try to import docx for .docx file support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ============================================================================
# CONFIGURATION PARAMETERS
# ============================================================================

# Default text file (can be overridden via command line)
TEXT_FILE = "C:/Users/Michael Kurian/Desktop/PhD Applications/genesis_original.txt.txt"

# Random seed for reproducibility
RANDOM_SEED = 42

# Window parameters - ADAPTIVE sizing
# The window size auto-adjusts to text length to produce ~110 windows
TARGET_WINDOWS = 110  # Adjust this to change granularity (higher = more windows)
# Note: Actual window size will be calculated as: total_tokens / (1 + (TARGET_WINDOWS - 1) / 2)

# ============================================================================
# MOTIF DICTIONARY - FELLOWSHIP OF THE RING (with phrases)
# ============================================================================

motif_dict = {
    # ========================================================================
    # TREE / AXIS MUNDI
    # Central axis connecting heaven and earth; knowledge and life
    # ========================================================================
    'Tree/Axis-Mundi': {
        'phrases': [
            'tree of life',           # Theological concept (appears 2x)
            'tree of knowledge',      # Central to Fall narrative (appears 2x)
            'tree of the knowledge',  # Full phrase variant
            'fruit of the tree',      # Action phrase
        ],
        'words': [
            'tree', 'trees', 'fruit', 'seed', 'yielding', 
            'herb', 'grass', 'plant', 'grow', 'grew', 
            'knowledge', 'bearing', 'green', 'ground', 
            'bring', 'brought', 'kind'
        ]
    },
    
    # ========================================================================
    # WATERS / FLOOD / SEA
    # Primordial waters, chaos, and life-giving fluid
    # ========================================================================
    'Waters/Sea': {
        'phrases': [
            'face of the deep',       # Primordial chaos (Genesis 1:2)
            'face of the waters',     # Spirit moving on waters
            'gathering together of the waters', # Separation act
            'divided the waters',     # Creation act
        ],
        'words': [
            'waters', 'water', 'seas', 'sea', 'river', 'deep', 
            'mist', 'rain', 'watered', 'gathered', 'parted', 
            'divided', 'face', 'pison', 'gihon', 'hiddekel', 
            'euphrates', 'havilah', 'ethiopia', 'assyria'
        ]
    },
    
    # ========================================================================
    # LIGHT
    # Divine light, celestial order, and temporal cycles
    # ========================================================================
    'Light': {
        'phrases': [
            'let there be light',     # Fiat lux - first creation (appears 1x)
            'there was light',        # Divine fulfillment
            'the evening and the morning', # Day formula (appears 6x)
            'divided the light',      # Separation act
        ],
        'words': [
            'light', 'lights', 'darkness', 'day', 'night', 
            'evening', 'morning', 'firmament', 'heaven', 'heavens', 
            'stars', 'rule', 'divide', 
            'signs', 'seasons', 'years', 'greater', 'lesser', 
            'first', 'second', 'third', 'fourth', 'fifth', 
            'sixth', 'seventh'
        ]
    },
    
    # ========================================================================
    # SHARPNESS / SWORD / THORN
    # Weapons, painful vegetation, barriers, and divine judgment
    # ========================================================================
    'Sharpness-Sword-Thorn': {
        'phrases': [
            'flaming sword',          # Guardian weapon at Eden's gate (Gen 3:24)
            'thorns also and thistles', # Curse on ground (Gen 3:18)
        ],
        'words': [
            'sword', 'flaming', 'thorns', 'thistles'
        ]
    },
    
    # ========================================================================
    # DEATH / MORTALITY
    # Death, dust, mortality, and return to earth
    # ========================================================================
    'Death/Mortality': {
        'phrases': [
            'shalt surely die',       # Divine warning (Gen 2:17, 3:4)
            'surely die',             # Serpent's contradiction variant
            'unto dust shalt thou return', # Mortality sentence (Gen 3:19)
            'dust shalt thou return', # Shorter variant
            'dust thou art',          # Mortality declaration
            'return unto the ground', # Return to earth
        ],
        'words': [
            'die', 'dust', 'return'
        ]
    },
    
    # ========================================================================
    # EATING / FORBIDDEN FRUIT
    # Consumption, eating, and the forbidden act
    # ========================================================================
    'Eating/Forbidden-Fruit': {
        'phrases': [
            'thou shalt not eat',     # Divine prohibition (Gen 2:17)
            'shalt not eat',          # Prohibition variant
            'ye shall not',           # Prohibition to Eve (Gen 3:3)
            'did eat',                # Fall action (Gen 3:6, 3:12)
            'i did eat',              # Adam's confession
            'she gave me',            # Adam blaming Eve
        ],
        'words': [
            'eat', 'eaten', 'eatest', 'fruit', 'food', 'meat'
        ]
    },
    
    # ========================================================================
    # CURSE / TOIL
    # Divine curse, sorrow, painful labor, and sweat
    # ========================================================================
    'Curse/Toil': {
        'phrases': [
            'cursed is the ground',   # Ground curse (Gen 3:17)
            'because thou hast',      # Judgment formula (Gen 3:14, 3:17)
            'in sorrow',              # Pain/sorrow formula (Gen 3:16, 3:17)
            'sweat of thy face',      # Toil consequence (Gen 3:19)
        ],
        'words': [
            'cursed', 'sorrow', 'sweat', 'till'
        ]
    },
    
    # ========================================================================
    # TRANSGRESSION / DECEPTION
    # Breaking commands, deception, and serpent's lies
    # ========================================================================
    'Transgression/Deception': {
        'phrases': [
            'ye shall be as gods',    # Serpent's promise (Gen 3:5)
            'shall be as gods',       # Variant
            'knowing good and evil',  # Temptation of knowledge (Gen 3:5)
            'the serpent beguiled',   # Eve's explanation (Gen 3:13)
            'serpent beguiled me',    # Confession variant
        ],
        'words': [
            'beguiled', 'subtil'
        ]
    },
    
    # ========================================================================
    # DESIRE / WISDOM
    # Visual temptation, desire for wisdom, and epistemological lust
    # ========================================================================
    'Desire/Wisdom': {
        'phrases': [
            'pleasant to the eyes',   # Visual temptation (Gen 3:6)
            'desired to make one wise', # Desire for wisdom (Gen 3:6)
            'to make one wise',       # Wisdom temptation variant
            'make one wise',          # Shorter variant
            'a tree to be desired',   # Desirability of tree (Gen 3:6)
            'good for food',          # Sensory appeal (Gen 3:6)
        ],
        'words': [
            'desired', 'desire', 'pleasant', 'wise'
        ]
    },
    
    # ========================================================================
    # FEAR / HIDING
    # Post-Fall fear, hiding from God, and psychological shame
    # ========================================================================
    'Fear/Hiding': {
        'phrases': [
            'i was afraid',           # Adam's fear (Gen 3:10)
            'i heard thy voice',      # Hearing with fear (Gen 3:10)
            'heard thy voice',        # Variant
            'i hid myself',           # Adam's hiding (Gen 3:10)
            'hid themselves',         # Couple hiding (Gen 3:8)
            'hid themselves from',    # Full hiding phrase (Gen 3:8)
        ],
        'words': [
            'afraid', 'hid'
        ]
    },
    
    # ========================================================================
    # EXILE / EXPULSION
    # Driving out from Eden, sending forth
    # ========================================================================
    'Exile/Expulsion': {
        'phrases': [
            'drove out the man',      # Expulsion act (Gen 3:24)
            'sent him forth',         # Sending from garden (Gen 3:23)
            'from the garden of eden', # Departure location
            'from the garden',        # Shorter variant
        ],
        'words': [
            'drove', 'sent', 'forth'
        ]
    },
    
    # ========================================================================
    # SERPENT / BEAST
    # Cunning beasts, chaos creatures, and animal life
    # ========================================================================
    'Serpent/Beast': {
        'phrases': [
            'beast of the field',     # Serpent's domain (appears 7x)
            'beast of the earth',     # Creation category
            'fowl of the air',        # Flying creatures (appears 5x)
            'fish of the sea',        # Aquatic creatures
            'living creature',        # General life
            'every living thing',     # Comprehensive life
        ],
        'words': [
            'serpent', 'beast', 'creature', 'creeping', 
            'creepeth', 'cattle', 'fowl', 'living', 'subtil', 
            'field', 'winged', 'moveth', 'whales', 'fly', 
            'fish', 'air', 'life', 'abundantly', 'multiply', 
            'enmity', 'bruise', 'belly', 'cursed', 'beguiled'
        ]
    },
    
    # ========================================================================
    # BRIDEGROOM-BRIDE
    # Marriage, union, human relationship and sexuality
    # ========================================================================
    'Bridegroom-Bride': {
        'phrases': [
            'male and female',        # Creation duality (appears 2x)
            'bone of my bones',       # Adam's recognition (appears 1x)
            'flesh of my flesh',      # Union phrase
            'one flesh',              # Marriage union (appears 1x)
            'help meet',              # Eve's role (appears 2x)
            'man and his wife',       # Marital pair
            'a living soul',          # Human essence (appears 1x)
        ],
        'words': [
            'man', 'woman', 'wife', 'husband', 'adam', 'eve', 
            'male', 'female', 'bone', 'bones', 'flesh', 'cleave', 
            'leave', 'father', 'mother', 'alone', 'help', 'meet', 
            'together', 'sleep', 'slept', 'ribs', 'rib', 'desire', 
            'conception', 'children', 'seed', 'living', 'hearkened', 
            
        ]
    },
    
    # ========================================================================
    # GARDEN / PARADISE
    # Sacred space, Eden, and geographical landmarks
    # ========================================================================
    'Garden/Paradise': {
        'phrases': [
            'garden of eden',         # Paradise location (appears 4x)
            'midst of the garden',    # Central location (appears 2x)
            'out of eden',            # Exile phrase
            'east of the garden',     # Post-exile location
        ],
        'words': [
            'garden', 'eden', 'pleasant', 'midst', 'east', 
            'eastward', 'planted', 'dress', 'keep', 'place', 
            'land', 'gold', 'bdellium', 'onyx', 'stone', 
            'compasseth', 'river', 'heads', 'parted', 'cherubims', 
            'way', 'whole', 'good', 'sight', 'food'
        ]
    },
    
    # ========================================================================
    # CLOTHING / NAKEDNESS
    # Shame, covering, and transformed consciousness
    # ========================================================================
    'Clothing/Nakedness': {
        'phrases': [
            'they were both naked',   # Pre-Fall description
            'the eyes of them both were opened', # Consciousness change
            'were not ashamed',       # Pre-Fall state
            'coats of skins',         # Divine covering
        ],
        'words': [
            'naked', 'clothed', 'coats', 'skins', 'aprons', 
            'fig', 'leaves', 'sewed', 'ashamed', 'opened', 
            'eyes', 'knew', 'hid', 'afraid', 'closed'
        ]
    },
    
    # ========================================================================
    # SACRED NAME / WORD
    # Divine speech, naming, and creative command
    # ========================================================================
    'Sacred-Name/Word': {
        'phrases': [
            'the lord god',           # Divine name (appears 20x)
            'and god said',           # Creation formula (appears 9x)
            'god said, let',          # Command structure (appears 8x)
            'and god saw',            # Divine approval (appears 7x)
            'god saw that',           # Evaluation phrase
            'it was good',            # Divine verdict (appears 7x)
            'in the beginning',       # Opening phrase (appears 1x)
            'image of god',           # Imago Dei (appears 3x)
            'spirit of god',          # Divine presence
            'breath of life',         # Life-giving act (appears 2x)
            'let us make',            # Divine counsel
        ],
        'words': [
            'god', 'lord', 'name', 'names', 'called', 'call', 
            'voice', 'commanded', 'saying', 'spirit', 'blessed', 
            'sanctified', 'created', 'made', 'rested', 'generations', 
            'breathed', 'formed', 'nostrils', 'breath', 'soul', 
            'image', 'likeness', 'dominion', 'replenish', 'subdue', 
            'behold', 'beginning', 'finished', 'host', 'work'
        ]
    },
    
    # ========================================================================
    # GOOD-EVIL / KNOWLEDGE
    # Moral knowledge, epistemological themes, and ethical consciousness
    # ========================================================================
    'Good-Evil/Knowledge': {
        'phrases': [
            'good and evil',          # Central moral concept (appears 4x)
            'knowledge of good',      # Epistemological phrase
        ],
        'words': [
            'good', 'evil', 'knowing', 'knowledge', 'wise', 
            'wisdom', 'eyes', 'opened', 'gods', 'die', 
            'death', 'surely', 'eat', 'eaten', 'touch'
        ]
    },
    
    # ========================================================================
    # COMMAND / OBEDIENCE
    # Divine imperatives, prohibitions, and speech formulas
    # ========================================================================
    'Command/Obedience': {
        'phrases': [
            'god commanded',          # Divine imperative
            'thou shalt not',         # Prohibition formula (appears 3x)
            'thou shalt surely',      # Emphasis formula
            'let there be',           # Creation command (appears 6x)
            'let them have',          # Dominion grant
            'said unto',              # Speech formula (appears 7x)
        ],
        'words': [
            'commanded', 'saying', 'thou', 'shalt', 'mayest', 
            'freely', 'whereof', 'shouldest',
        ]
    },
    
    # ========================================================================
    # TIME / ORDER
    # Temporal structure, cosmological sequence, and creation days
    # ========================================================================
    'Time/Order': {
        'phrases': [
            'in the beginning',       # Temporal origin
            'the first day',          # Creation sequence (appears 6x)
            'the second day',         # Day 2
            'the third day',          # Day 3
            'the fourth day',         # Day 4
            'the fifth day',          # Day 5
            'the sixth day',          # Day 6
            'the seventh day',        # Sabbath (appears 3x)
            'evening and morning',    # Day formula (appears 6x)
        ],
        'words': [
            'day', 'days', 'beginning', 'first', 'second', 
            'third', 'fourth', 'fifth', 'sixth', 'seventh', 
            'evening', 'morning', 'seasons', 'years', 
            'generations', 'finished'
        ]
    }
}

# ============================================================================
# PHRASE PRE-PROCESSING
# ============================================================================

def extract_all_phrases(motif_dict):
    """
    Extract all multi-word phrases from motif dictionary.
    
    Returns:
        list: All phrases, sorted by length (longest first) to avoid subset conflicts
    """
    all_phrases = []
    for category_data in motif_dict.values():
        if isinstance(category_data, dict) and 'phrases' in category_data:
            all_phrases.extend(category_data['phrases'])
    
    # Sort by word count (longest first) to match longer phrases first
    all_phrases.sort(key=lambda x: len(x.split()), reverse=True)
    return all_phrases


def merge_phrases_in_text(text, phrases):
    """
    Pre-process text to merge multi-word phrases into single tokens.
    
    Example: "the One Ring was" → "the one_ring was"
    
    Args:
        text (str): Raw input text
        phrases (list): List of multi-word phrases to merge
        
    Returns:
        str: Text with phrases merged using underscores
    """
    # Work on lowercase for matching, preserve original for context
    text_lower = text.lower()
    
    # Track replacements to avoid overlapping substitutions
    replacements = []
    
    for phrase in phrases:
        phrase_lower = phrase.lower()
        phrase_merged = phrase_lower.replace(' ', '_')
        
        # Use word boundaries to ensure whole phrase matching
        # Pattern: \b word1 \s+ word2 \s+ word3 \b
        pattern_parts = phrase_lower.split()
        pattern = r'\b' + r'\s+'.join(re.escape(p) for p in pattern_parts) + r'\b'
        
        # Find all matches
        for match in re.finditer(pattern, text_lower):
            start, end = match.span()
            replacements.append((start, end, phrase_merged))
    
    # Sort replacements by start position (reverse) to replace from end to beginning
    replacements.sort(key=lambda x: x[0], reverse=True)
    
    # Apply replacements
    text_merged = text_lower
    for start, end, replacement in replacements:
        text_merged = text_merged[:start] + replacement + text_merged[end:]
    
    return text_merged


# ============================================================================
# FILE LOADING FUNCTIONS
# ============================================================================

def load_text_file(filepath):
    """
    Load text from .txt or .docx file with smart encoding detection.
    
    Args:
        filepath: Path to text file (.txt) or Word document (.docx)
        
    Returns:
        str: Text content
    """
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()
    
    if ext == '.docx':
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for .docx files. Install with: pip install python-docx")
        doc = Document(filepath)
        text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        print(f"✓ Loaded .docx file: {len(doc.paragraphs)} paragraphs")
        return text
    elif ext == '.txt':
        # Try common encodings
        for encoding in ['utf-8', 'utf-8-sig', 'cp1252', 'latin-1', 'iso-8859-1']:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    text = f.read()
                print(f"✓ Loaded with {encoding} encoding")
                return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        # Last resort: utf-8 with error handling
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        print(f"⚠ Loaded with utf-8 (some characters replaced)")
        return text
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .txt or .docx")


# ============================================================================
# TOKENIZATION (with merged phrases)
# ============================================================================

def tokenize_text(text):
    """
    Tokenize text using word boundaries and lowercase normalization.
    Handles merged phrases (with underscores) as single tokens.
    
    Args:
        text (str): Raw text input (already has phrases merged)
        
    Returns:
        list: List of lowercase tokens (including merged phrase tokens)
    """
    # Remove punctuation except underscores (phrase markers) and apostrophes
    text = re.sub(r'[^a-zA-Z\s\'_]', ' ', text.lower())
    # Match words and underscore-connected phrases
    tokens = re.findall(r'\b[\w_]+(?:\'[\w_]+)?\b', text)
    return tokens


# ============================================================================
# MOTIF COUNTING
# ============================================================================

def count_motifs_in_window(window_tokens, motif_dict):
    """
    Count motif occurrences in a window.
    Handles both single words and merged phrases (with underscores).
    
    Args:
        window_tokens (list): List of tokens in current window
        motif_dict (dict): Dictionary with 'phrases' and/or 'words' keys
        
    Returns:
        dict: {category: count} for each motif category
    """
    counts = {category: 0 for category in motif_dict.keys()}
    window_set = set(window_tokens)
    
    for category, category_data in motif_dict.items():
        # Handle new dictionary structure
        if isinstance(category_data, dict):
            # Count merged phrases (they appear as underscore-connected tokens)
            if 'phrases' in category_data:
                for phrase in category_data['phrases']:
                    merged_phrase = phrase.lower().replace(' ', '_')
                    if merged_phrase in window_set:
                        counts[category] += window_tokens.count(merged_phrase)
            
            # Count single words
            if 'words' in category_data:
                for word in category_data['words']:
                    word_lower = word.lower()
                    if word_lower in window_set:
                        counts[category] += window_tokens.count(word_lower)
        else:
            # Legacy format: list of words
            for word in category_data:
                word_lower = word.lower()
                if word_lower in window_set:
                    counts[category] += window_tokens.count(word_lower)
    
    return counts


# ============================================================================
# BASELINE CALCULATION
# ============================================================================

def calculate_global_baseline(tokens, motif_dict):
    """
    Calculate global baseline as proportion of each motif in full text.
    
    Formula: π_k = (total count of motif k) / (total semantic tokens)
    
    Args:
        tokens (list): All tokens in text (with merged phrases)
        motif_dict (dict): Motif categories
        
    Returns:
        dict: {category: proportion} baseline values
    """
    N = len(tokens)
    baseline = {}
    
    for category, category_data in motif_dict.items():
        total_count = 0
        
        # Handle new dictionary structure
        if isinstance(category_data, dict):
            # Count merged phrases
            if 'phrases' in category_data:
                for phrase in category_data['phrases']:
                    merged_phrase = phrase.lower().replace(' ', '_')
                    total_count += tokens.count(merged_phrase)
            
            # Count single words
            if 'words' in category_data:
                for word in category_data['words']:
                    total_count += tokens.count(word.lower())
        else:
            # Legacy format
            for word in category_data:
                total_count += tokens.count(word.lower())
        
        baseline[category] = total_count / N if N > 0 else 0
    
    return baseline


# ============================================================================
# SHANNON ENTROPY
# ============================================================================

def calculate_shannon_entropy(window_tokens):
    """
    Calculate Shannon entropy (H) for a window.
    
    Formula: H = -Σ p(x) * log₂(p(x))
    
    Args:
        window_tokens (list): Tokens in current window
        
    Returns:
        float: Shannon entropy in bits per semantic token
    """
    if len(window_tokens) == 0:
        return 0.0
    
    word_counts = Counter(window_tokens)
    total = len(window_tokens)
    
    H = 0.0
    for count in word_counts.values():
        p = count / total
        if p > 0:
            H -= p * np.log2(p)
    
    return H


# ============================================================================
# KL DIVERGENCE (SIGMA)
# ============================================================================

def calculate_sigma_kl(observed, baseline_proportions, window_size):
    """
    Calculate Sigma (Σ) using true KL divergence formula.
    
    Formula: Σ_KL = Σ p_k * log₂(p_k / π_k)
    
    Where:
        p_k = observed proportion of motif k in window
        π_k = baseline proportion of motif k
    
    Args:
        observed (dict): {category: count} in current window
        baseline_proportions (dict): {category: proportion} baseline
        window_size (int): Size of window for normalization
        
    Returns:
        float: KL divergence in bits per semantic token
    """
    sigma_kl = 0.0
    
    for category in observed.keys():
        obs_count = observed[category]
        p_k = obs_count / window_size
        pi_k = baseline_proportions[category]
        
        if p_k > 0 and pi_k > 0:
            sigma_kl += p_k * np.log2(p_k / pi_k)
    
    return sigma_kl


# ============================================================================
# MAIN ANALYSIS PIPELINE
# ============================================================================

def run_se_analysis(text_path, motif_dict):
    """
    Execute complete Symbolic Entropy analysis pipeline.
    
    Args:
        text_path (str): Path to input text file
        motif_dict (dict): Motif category definitions
        
    Returns:
        tuple: (results_df, raw_densities, kl_contributions, motif_dict, 
                baseline, window_size, total_tokens, tokens)
    """
    print(f"\n{'='*70}")
    print(f"SYMBOLIC ENTROPY (SE) ANALYSIS - Multi-word Phrase Support")
    print(f"{'='*70}\n")
    
    # Validate file exists
    if not os.path.exists(text_path):
        raise FileNotFoundError(f"Text file not found: {text_path}")
    
    print(f"Loading text: {os.path.basename(text_path)}")
    text = load_text_file(text_path)
    
    # Extract and merge phrases
    print("\nPre-processing multi-word phrases...")
    all_phrases = extract_all_phrases(motif_dict)
    if all_phrases:
        print(f"  Found {len(all_phrases)} phrases to merge")
        print(f"  Examples: {', '.join(all_phrases[:3])}")
        text = merge_phrases_in_text(text, all_phrases)
    else:
        print("  No phrases found (word-only analysis)")
    
    # Tokenization
    print("\nTokenizing...")
    tokens = tokenize_text(text)
    total_tokens = len(tokens)
    print(f"✓ Total semantic tokens: {total_tokens:,}")
    
    # Display motif structure
    n_categories = len(motif_dict)
    print(f"✓ Motif categories: {n_categories}")
    
    # Calculate global baseline
    print("\nCalculating global baseline...")
    baseline = calculate_global_baseline(tokens, motif_dict)
    
    # ADAPTIVE WINDOW SIZING - targets ~110 windows regardless of text length
    # Formula maintains 50% overlap: window_size = total_tokens / (1 + (target - 1) / 2)
    window_size = int(total_tokens / (1 + (TARGET_WINDOWS - 1) / 2))
    step_size = window_size // 2  # MANDATORY: 50% overlap
    n_windows = (total_tokens - window_size) // step_size + 1
    
    print(f"\n{'='*70}")
    print(f"ADAPTIVE WINDOW PARAMETERS")
    print(f"{'='*70}")
    print(f"Target windows: {TARGET_WINDOWS}")
    print(f"Calculated window size: {window_size} tokens")
    print(f"Step size: {step_size} tokens (50% overlap)")
    print(f"Actual windows: {n_windows}")
    print(f"{'='*70}")
    
    # Sliding window analysis
    print("\nAnalyzing windows...")
    results = []
    raw_densities = []
    kl_contributions = []
    
    n_windows = (total_tokens - window_size) // step_size + 1
    
    for i in range(0, total_tokens - window_size + 1, step_size):
        window_tokens = tokens[i:i + window_size]
        
        # Count motifs
        observed = count_motifs_in_window(window_tokens, motif_dict)
        
        # Calculate H
        H = calculate_shannon_entropy(window_tokens)
        
        # Calculate Σ
        Sigma = calculate_sigma_kl(observed, baseline, window_size)
        
        # Store results
        results.append({
            'window_index': len(results),
            'start_token': i,
            'end_token': i + window_size,
            'H': H,
            'Sigma': Sigma,
            'SE': H + Sigma
        })
        
        # Store raw densities (for heatmap)
        raw_densities.append([observed[cat] / window_size for cat in motif_dict.keys()])
        
        # Store KL contributions (for heatmap)
        kl_row = []
        for category in motif_dict.keys():
            p_k = observed[category] / window_size
            pi_k = baseline[category]
            if p_k > 0 and pi_k > 0:
                kl_row.append(p_k * np.log2(p_k / pi_k))
            else:
                kl_row.append(0.0)
        kl_contributions.append(kl_row)
    
    print(f"✓ Analyzed {len(results)} windows")
    
    # Convert to DataFrame
    results_df = pd.DataFrame(results)
    raw_densities = np.array(raw_densities)
    kl_contributions = np.array(kl_contributions)
    
    return (results_df, raw_densities, kl_contributions, motif_dict, 
            baseline, window_size, total_tokens, tokens)


# ============================================================================
# VISUALIZATION - DUAL HEATMAP
# ============================================================================

def plot_dual_heatmap(results_df, raw_densities, kl_contributions, 
                      motif_dict, output_prefix):
    """
    Generate dual heatmap: raw motif density + KL divergence contributions.
    Includes overlaid line plots of H (cyan) and Σ (white) on the KL heatmap.
    """
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(18, 8))
    
    categories = list(motif_dict.keys())
    n_categories = len(categories)
    window_indices = results_df['window_index'].values
    
    # Normalize H and Σ for overlay (scale to heatmap Y-axis range)
    H_values = results_df['H'].values
    sigma_values = results_df['Sigma'].values
    
    # Min-max normalization to [0, n_categories-1] range
    H_min, H_max = H_values.min(), H_values.max()
    sigma_min, sigma_max = sigma_values.min(), sigma_values.max()
    
    H_scaled = (n_categories - 1) * (1 - (H_values - H_min) / (H_max - H_min + 1e-10))
    sigma_scaled = (n_categories - 1) * (1 - (sigma_values - sigma_min) / (sigma_max - sigma_min + 1e-10))
    
    # LEFT: Raw density heatmap
    raw_data = raw_densities.T
    im1 = ax1.imshow(raw_data, aspect='auto', cmap='plasma', 
                     interpolation='nearest', origin='lower')
    ax1.set_ylim(-0.5, n_categories - 0.5)
    ax1.set_xlabel('Window Index', fontsize=12)
    ax1.set_ylabel('Motif Category', fontsize=12)
    ax1.set_yticks(range(n_categories))
    ax1.set_yticklabels(categories, fontsize=9)
    ax1.set_title('Method 1: RAW DENSITY\n(Simple frequency counting)', 
                  fontsize=13, fontweight='bold')
    
    cbar1 = plt.colorbar(im1, ax=ax1)
    cbar1.set_label('Proportion of Window', fontsize=10)
    
    # RIGHT: KL Divergence with line overlays
    kl_data = kl_contributions.T
    im2 = ax2.imshow(kl_data, aspect='auto', cmap='plasma', 
                     interpolation='nearest', origin='lower')
    
    # Overlay H and Σ lines (thin and transparent to not block heatmap)
    ax2.plot(window_indices, sigma_scaled, color='white', linewidth=1.0, 
             alpha=0.4, label='Σ (white)')
    ax2.plot(window_indices, H_scaled, color='cyan', linewidth=1.0, 
             alpha=0.4, label='H (cyan)')
    
    ax2.set_ylim(-0.5, n_categories - 0.5)
    ax2.set_xlabel('Window Index', fontsize=12)
    ax2.set_ylabel('Motif Category', fontsize=12)
    ax2.set_yticks(range(n_categories))
    ax2.set_yticklabels(categories, fontsize=9)
    ax2.set_title('Method 2: KL DIVERGENCE (Σ_KL)\n(Structural Surprise - Where motifs cluster)', 
                  fontsize=13, fontweight='bold')
    
    cbar2 = plt.colorbar(im2, ax=ax2)
    cbar2.set_label('KL Contribution (bits/token)', fontsize=10)
    
    # Add legend for line overlays
    ax2.legend(loc='upper right', fontsize=9, framealpha=0.8)
    
    fig.suptitle('Symbolic Entropy Analysis - Dual Method Visualization', 
                 fontsize=16, fontweight='bold', y=0.98)
    
    plt.tight_layout()
    filename = f'{output_prefix}_se_heatmap.png'
    plt.savefig(filename, dpi=300, bbox_inches='tight')
    print(f"\n✓ Saved: {filename}")
    plt.show()
    plt.close()


# ============================================================================
# VISUALIZATION - TIME SERIES
# ============================================================================

def plot_timeseries(results_df, output_prefix):
    """
    Generate separate line plots for H, Σ, and SE.
    """
    fig, axes = plt.subplots(3, 1, figsize=(14, 10), sharex=True)
    
    # H plot
    axes[0].plot(results_df['window_index'], results_df['H'], 
                 color='steelblue', linewidth=1.5)
    axes[0].set_ylabel('H (bits/token)', fontsize=11)
    axes[0].set_title('Shannon Entropy (H) - Lexical Diversity', 
                      fontsize=12, fontweight='bold')
    axes[0].grid(True, alpha=0.3)
    
    # Σ plot
    axes[1].plot(results_df['window_index'], results_df['Sigma'], 
                 color='crimson', linewidth=1.5)
    axes[1].set_ylabel('Σ (bits/token)', fontsize=11)
    axes[1].set_title('Sigma (Σ) - Archetypal Motif Concentration', 
                      fontsize=12, fontweight='bold')
    axes[1].grid(True, alpha=0.3)
    
    # SE plot
    axes[2].plot(results_df['window_index'], results_df['SE'], 
                 color='darkgreen', linewidth=1.5)
    axes[2].set_ylabel('SE (bits/token)', fontsize=11)
    axes[2].set_xlabel('Window Index', fontsize=11)
    axes[2].set_title('Symbolic Entropy (SE = H + Σ)', 
                      fontsize=12, fontweight='bold')
    axes[2].grid(True, alpha=0.3)
    
    plt.tight_layout()
    filename = f'{output_prefix}_se_timeseries.png'
    plt.savefig(filename, dpi=300, bbox_inches='tight')
    print(f"✓ Saved: {filename}")
    plt.show()
    plt.close()


# ============================================================================
# PEAK AND VALLEY ANALYSIS
# ============================================================================

def extract_window_text(window_idx, tokens, window_size, step_size, context_words=50):
    """
    Extract the text corresponding to a window, with optional context.
    
    Args:
        window_idx: Index of the window
        tokens: Full token list
        window_size: Size of each window
        step_size: Step between windows
        context_words: Additional words before/after for context
        
    Returns:
        dict with 'window_text', 'start_position', 'end_position'
    """
    start = window_idx * step_size
    end = start + window_size
    
    # Add context
    context_start = max(0, start - context_words)
    context_end = min(len(tokens), end + context_words)
    
    # Extract tokens (unmerg underscores for readability)
    window_tokens = tokens[context_start:context_end]
    text = ' '.join(window_tokens).replace('_', ' ')
    
    return {
        'window_text': text,
        'start_position': start,
        'end_position': end
    }


def analyze_window_motifs(window_tokens, motif_dict):
    """
    Analyze which motifs are most present in a window.
    
    Returns:
        dict: {category: count} sorted by count (descending)
    """
    counts = count_motifs_in_window(window_tokens, motif_dict)
    # Sort by count
    sorted_counts = dict(sorted(counts.items(), key=lambda x: x[1], reverse=True))
    # Filter out zero counts
    return {k: v for k, v in sorted_counts.items() if v > 0}


def plot_peaks_and_valleys(results_df, tokens, window_size, step_size,
                           motif_dict, output_prefix):
    """
    Create visualization showing Σ line graph with annotated peaks/valleys,
    plus text excerpts from key moments.
    """
    print(f"\n{'='*70}")
    print(f"GENERATING PEAK AND VALLEY ANALYSIS")
    print(f"{'='*70}\n")
    
    # Find peaks and valleys using scipy
    sigma_values = results_df['Sigma'].values
    
    # Find peaks (local maxima)
    peak_indices, _ = find_peaks(sigma_values, distance=5, prominence=0.001)
    if len(peak_indices) > 0:
        peak_values = sigma_values[peak_indices]
        # Get top 3
        top_peak_idx = np.argsort(peak_values)[-3:][::-1]
        peaks = [(peak_indices[i], peak_values[i]) for i in top_peak_idx]
    else:
        # Fallback: just use top 3 values
        top_3_idx = np.argsort(sigma_values)[-3:][::-1]
        peaks = [(idx, sigma_values[idx]) for idx in top_3_idx]
    
    # Find valleys (local minima)
    valley_indices, _ = find_peaks(-sigma_values, distance=5, prominence=0.001)
    if len(valley_indices) > 0:
        valley_values = sigma_values[valley_indices]
        # Get bottom 3
        bottom_valley_idx = np.argsort(valley_values)[:3]
        valleys = [(valley_indices[i], valley_values[i]) for i in bottom_valley_idx]
    else:
        # Fallback: just use bottom 3 values
        bottom_3_idx = np.argsort(sigma_values)[:3]
        valleys = [(idx, sigma_values[idx]) for idx in bottom_3_idx]
    
    print(f"Top 3 Peaks (Highest Σ):")
    for rank, (idx, val) in enumerate(peaks, 1):
        print(f"  Peak {rank}: Window {idx}, Σ = {val:.6f}")
    
    print(f"\nTop 3 Valleys (Lowest Σ):")
    for rank, (idx, val) in enumerate(valleys, 1):
        print(f"  Valley {rank}: Window {idx}, Σ = {val:.6f}")
    
    # Create figure: line graph at top + text excerpts below
    fig = plt.figure(figsize=(16, 14))
    gs = fig.add_gridspec(4, 2, height_ratios=[3, 1, 1, 1], hspace=0.5, wspace=0.4)
    
    # MAIN PLOT: Σ over time with highlighted peaks and valleys
    ax_main = fig.add_subplot(gs[0, :])
    
    window_indices = results_df['window_index'].values
    
    # Plot Σ line
    ax_main.plot(window_indices, sigma_values, color='crimson', 
                 linewidth=2, label='Σ (Motif Concentration)', alpha=0.7)
    
    # Mark peaks with upward arrows
    for rank, (idx, val) in enumerate(peaks, 1):
        ax_main.scatter(idx, val, color='gold', s=300, marker='^', 
                       edgecolors='black', linewidths=2, zorder=5)
        ax_main.annotate(f'Peak {rank}', xy=(idx, val), 
                        xytext=(0, 20), textcoords='offset points',
                        ha='center', fontsize=11, fontweight='bold',
                        bbox=dict(boxstyle='round,pad=0.5', facecolor='gold', alpha=0.8),
                        arrowprops=dict(arrowstyle='->', lw=2))
    
    # Mark valleys with downward arrows
    for rank, (idx, val) in enumerate(valleys, 1):
        ax_main.scatter(idx, val, color='lightblue', s=300, marker='v', 
                       edgecolors='black', linewidths=2, zorder=5)
        ax_main.annotate(f'Valley {rank}', xy=(idx, val), 
                        xytext=(0, -20), textcoords='offset points',
                        ha='center', fontsize=11, fontweight='bold',
                        bbox=dict(boxstyle='round,pad=0.5', facecolor='lightblue', alpha=0.8),
                        arrowprops=dict(arrowstyle='->', lw=2))
    
    ax_main.set_xlabel('Window Index', fontsize=12, fontweight='bold')
    ax_main.set_ylabel('Σ (bits/token)', fontsize=12, fontweight='bold')
    ax_main.set_title('Symbolic Entropy: Key Moments (Top 3 Peaks & Valleys)', 
                     fontsize=14, fontweight='bold')
    ax_main.grid(True, alpha=0.3)
    ax_main.legend(loc='upper right', fontsize=10)
    
    # TEXT EXCERPT PANELS - 3x2 grid
    excerpt_axes = [
        (fig.add_subplot(gs[1, 0]), 'Peak 1', peaks[0] if len(peaks) > 0 else None, 'gold'),
        (fig.add_subplot(gs[1, 1]), 'Peak 2', peaks[1] if len(peaks) > 1 else None, 'gold'),
        (fig.add_subplot(gs[2, 0]), 'Peak 3', peaks[2] if len(peaks) > 2 else None, 'gold'),
        (fig.add_subplot(gs[2, 1]), 'Valley 1', valleys[0] if len(valleys) > 0 else None, 'lightblue'),
        (fig.add_subplot(gs[3, 0]), 'Valley 2', valleys[1] if len(valleys) > 1 else None, 'lightblue'),
        (fig.add_subplot(gs[3, 1]), 'Valley 3', valleys[2] if len(valleys) > 2 else None, 'lightblue'),
    ]
    
    for ax, label, point_data, bgcolor in excerpt_axes:
        ax.axis('off')
        
        if point_data is None:
            continue
            
        idx, val = point_data
        
        # Get text excerpt
        text_data = extract_window_text(idx, tokens, window_size, step_size, context_words=50)
        
        # Analyze motifs
        start = text_data['start_position']
        end = text_data['end_position']
        window_tokens = tokens[start:end]
        top_motifs = analyze_window_motifs(window_tokens, motif_dict)
        top_3_motifs = list(top_motifs.items())[:3]
        motif_str = ', '.join([f"{cat}: {cnt}" for cat, cnt in top_3_motifs]) if top_3_motifs else "None"
        
        # Format text excerpt
        excerpt = text_data['window_text']
        if len(excerpt) > 250:
            excerpt = excerpt[:250] + '...'
        
        # Word wrap for display
        words = excerpt.split()
        lines = []
        current_line = []
        current_length = 0
        for word in words:
            if current_length + len(word) + 1 > 40:
                lines.append(' '.join(current_line))
                current_line = [word]
                current_length = len(word)
            else:
                current_line.append(word)
                current_length += len(word) + 1
        if current_line:
            lines.append(' '.join(current_line))
        wrapped_excerpt = '\n'.join(lines[:5])
        
        # Display
        title_text = f"{label} (Win {idx}, Σ={val:.5f})"
        motif_text = f"Motifs: {motif_str}"
        
        ax.text(0.5, 0.95, title_text, 
               transform=ax.transAxes, fontsize=10, fontweight='bold',
               ha='center', va='top',
               bbox=dict(boxstyle='round,pad=0.4', facecolor=bgcolor, alpha=0.6))
        
        ax.text(0.5, 0.75, motif_text, 
               transform=ax.transAxes, fontsize=8,
               ha='center', va='top', style='italic',
               bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.7))
        
        ax.text(0.5, 0.50, wrapped_excerpt, 
               transform=ax.transAxes, fontsize=7,
               ha='center', va='top', family='monospace',
               bbox=dict(boxstyle='round,pad=0.4', facecolor='lightyellow', 
                        alpha=0.8, edgecolor='gray', linewidth=0.5))
    
    plt.tight_layout()
    filename = f'{output_prefix}_peaks_valleys.png'
    plt.savefig(filename, dpi=300, bbox_inches='tight')
    print(f"✓ Saved: {filename}")
    plt.show()
    plt.close()
    
    # Export detailed CSV
    export_peaks_valleys_csv(peaks, valleys, tokens, window_size, step_size, 
                             motif_dict, output_prefix)


def export_peaks_valleys_csv(peaks, valleys, tokens, window_size, step_size,
                             motif_dict, output_prefix):
    """
    Export peak and valley text excerpts to CSV for detailed analysis.
    """
    data = []
    
    # Add peaks
    for rank, (idx, val) in enumerate(peaks, 1):
        text_data = extract_window_text(idx, tokens, window_size, step_size, 
                                       context_words=100)
        start = text_data['start_position']
        end = text_data['end_position']
        window_tokens = tokens[start:end]
        motifs = analyze_window_motifs(window_tokens, motif_dict)
        
        data.append({
            'type': 'Peak',
            'rank': rank,
            'window_idx': idx,
            'sigma_value': val,
            'excerpt': text_data['window_text'][:500],
            'top_motif_1': list(motifs.keys())[0] if len(motifs) > 0 else '',
            'top_motif_1_count': list(motifs.values())[0] if len(motifs) > 0 else 0,
            'top_motif_2': list(motifs.keys())[1] if len(motifs) > 1 else '',
            'top_motif_2_count': list(motifs.values())[1] if len(motifs) > 1 else 0,
            'top_motif_3': list(motifs.keys())[2] if len(motifs) > 2 else '',
            'top_motif_3_count': list(motifs.values())[2] if len(motifs) > 2 else 0,
        })
    
    # Add valleys
    for rank, (idx, val) in enumerate(valleys, 1):
        text_data = extract_window_text(idx, tokens, window_size, step_size, 
                                       context_words=100)
        start = text_data['start_position']
        end = text_data['end_position']
        window_tokens = tokens[start:end]
        motifs = analyze_window_motifs(window_tokens, motif_dict)
        
        data.append({
            'type': 'Valley',
            'rank': rank,
            'window_idx': idx,
            'sigma_value': val,
            'excerpt': text_data['window_text'][:500],
            'top_motif_1': list(motifs.keys())[0] if len(motifs) > 0 else '',
            'top_motif_1_count': list(motifs.values())[0] if len(motifs) > 0 else 0,
            'top_motif_2': list(motifs.keys())[1] if len(motifs) > 1 else '',
            'top_motif_2_count': list(motifs.values())[1] if len(motifs) > 1 else 0,
            'top_motif_3': list(motifs.keys())[2] if len(motifs) > 2 else '',
            'top_motif_3_count': list(motifs.values())[2] if len(motifs) > 2 else 0,
        })
    
    df = pd.DataFrame(data)
    filename = f'{output_prefix}_peaks_valleys_text.csv'
    df.to_csv(filename, index=False)
    print(f"✓ Saved: {filename}")


# ============================================================================
# PUBLICATION STATISTICS
# ============================================================================

def print_publication_statistics(results_df, text_path, total_tokens, 
                                 window_size, n_windows, n_phrases):
    """
    Print comprehensive statistics for publication.
    """
    print(f"\n{'='*70}")
    print(f"PUBLICATION-READY STATISTICS")
    print(f"{'='*70}")
    print(f"Text: {os.path.basename(text_path)}")
    print(f"Total semantic tokens: {total_tokens:,}")
    print(f"  (includes {n_phrases} merged multi-word phrases)")
    print(f"Window size: {window_size} tokens")
    print(f"Number of windows: {n_windows}")
    print(f"")
    print(f"Shannon Entropy (H):")
    print(f"  Mean: {results_df['H'].mean():.4f} ± {results_df['H'].std():.4f} bits/token")
    print(f"  Range: [{results_df['H'].min():.4f}, {results_df['H'].max():.4f}]")
    print(f"")
    print(f"Sigma (Σ):")
    print(f"  Mean: {results_df['Sigma'].mean():.6f} ± {results_df['Sigma'].std():.6f} bits/token")
    print(f"  Range: [{results_df['Sigma'].min():.6f}, {results_df['Sigma'].max():.6f}]")
    print(f"")
    print(f"Symbolic Entropy (SE = H + Σ):")
    print(f"  Mean: {results_df['SE'].mean():.4f} ± {results_df['SE'].std():.4f} bits/token")
    print(f"  Range: [{results_df['SE'].min():.4f}, {results_df['SE'].max():.4f}]")
    print(f"")
    print(f"Note: 'bits/token' refers to bits per SEMANTIC TOKEN")
    print(f"      Multi-word phrases are merged into single tokens")
    print(f"{'='*70}\n")


# ============================================================================
# MAIN EXECUTION
# ============================================================================

if __name__ == "__main__":
    # Set random seed for reproducibility
    random.seed(RANDOM_SEED)
    np.random.seed(RANDOM_SEED)
    
    # Handle command-line argument
    if len(sys.argv) > 1:
        TEXT_FILE = sys.argv[1]
        print(f"Using command-line specified file: {TEXT_FILE}")
    
    # Check file exists
    if not os.path.exists(TEXT_FILE):
        print(f"ERROR: File not found: {TEXT_FILE}")
        print(f"Usage: python {sys.argv[0]} [path_to_text_file]")
        sys.exit(1)
    
    # Generate output prefix from filename
    output_prefix = os.path.splitext(os.path.basename(TEXT_FILE))[0]
    
    # Run main analysis
    try:
        (results_df, raw_densities, kl_contributions, motif_dict_used, 
         baseline, window_size, total_tokens, tokens) = run_se_analysis(TEXT_FILE, motif_dict)
    except Exception as e:
        print(f"ERROR during analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    
    n_windows = len(results_df)
    
    # Count merged phrases for statistics
    n_phrases = sum(1 for token in tokens if '_' in token)
    
    # Print publication statistics
    print_publication_statistics(results_df, TEXT_FILE, total_tokens, 
                                 window_size, n_windows, n_phrases)
    
    # Generate visualizations
    plot_dual_heatmap(results_df, raw_densities, kl_contributions, 
                     motif_dict_used, output_prefix)
    
    plot_timeseries(results_df, output_prefix)
    
    # Generate peak and valley visualization with text excerpts
    step_size = window_size // 2  # 50% overlap
    plot_peaks_and_valleys(results_df, tokens, window_size, step_size,
                          motif_dict_used, output_prefix)
    
    # Save results
    csv_filename = f'{output_prefix}_se_results.csv'
    results_df.to_csv(csv_filename, index=False)
    print(f"\n✓ Results saved: {csv_filename}")
    
    # Final summary
    print(f"\n{'='*70}")
    print(f"ANALYSIS COMPLETE")
    print(f"{'='*70}")
    print(f"Outputs generated:")
    print(f"  - {output_prefix}_se_heatmap.png")
    print(f"  - {output_prefix}_se_timeseries.png")
    print(f"  - {output_prefix}_peaks_valleys.png")
    print(f"  - {output_prefix}_se_results.csv")
    print(f"  - {output_prefix}_peaks_valleys_text.csv")
    print(f"\nFor publication, cite:")
    print(f"  Kurian, M. (2025). Symbolic Entropy: A Mathematical Framework")
    print(f"  for Quantifying Meaning Density in Text.")
    print(f"{'='*70}\n")
